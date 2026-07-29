"""Build the Phase 8 Word report from retained, machine-readable evidence."""

from __future__ import annotations

import argparse
import csv
import json
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "paper" / "report-content.md"
SUMMARY_PATH = ROOT / "results" / "tables" / "phase-07-summary.json"
STRATA_PATH = ROOT / "results" / "tables" / "phase-07-stratified-outcomes.csv"
RISK_PATH = ROOT / "results" / "tables" / "phase-07-risk-register.jsonl"
CATALOG_PATH = ROOT / "configs" / "phase-05-scenario-catalog.v1.1.0.json"
CLAIMS_PATH = ROOT / "paper" / "claim-evidence.json"
FIGURE_PATH = ROOT / "paper" / "generated" / "phase-08-category-outcomes.png"

INK = RGBColor(0x10, 0x18, 0x28)
MUTED = RGBColor(0x47, 0x54, 0x67)
BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
SUCCESS = (180, 35, 24)
FAILURE = (20, 122, 82)
AMBIGUOUS = (181, 71, 8)
ERROR = (102, 112, 133)


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as stream:
        return list(csv.DictReader(stream))


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def outcome(summary: dict[str, Any], population: str) -> dict[str, Any]:
    for item in summary["outcomes"]:
        if item["population"] == population:
            return item
    raise KeyError(population)


def category_rows() -> list[dict[str, str]]:
    return [
        row
        for row in read_csv(STRATA_PATH)
        if row["population"] == "adversarial" and row["dimension"] == "category"
    ]


def environment_tokens() -> dict[str, str]:
    system = read_json(ROOT / "environment" / "system-info.json")
    models = read_json(ROOT / "environment" / "ollama-models.json")
    python_lines = dict(
        line.split("=", 1)
        for line in (ROOT / "environment" / "python-environment.txt")
        .read_text(encoding="utf-8")
        .splitlines()
        if "=" in line
    )
    image_digest = dict(
        line.split("=", 1)
        for line in (ROOT / "environment" / "pwnzzai-image-digest.txt")
        .read_text(encoding="utf-8")
        .splitlines()
        if "=" in line
    )
    principal = models["principal_model"]
    return {
        "GARAK_VERSION": python_lines["garak_version"],
        "PWNZZAI_COMMIT": (ROOT / "environment" / "pwnzzai-commit.txt")
        .read_text(encoding="utf-8")
        .strip(),
        "PWNZZAI_IMAGE_DIGEST": image_digest["registry_manifest_digest"],
        "OLLAMA_VERSION": models["ollama_version"],
        "MODEL_TAG": principal["tag"],
        "MODEL_DIGEST": principal["digest"],
        "MODEL_PARAMETER_SIZE": principal["parameter_size"],
        "MODEL_QUANTIZATION": principal["quantization"],
        "WINDOWS_BUILD": system["host"]["build"],
        "DOCKER_DESKTOP_VERSION": system["docker"]["desktop_version"],
        "DOCKER_ENGINE_VERSION": system["docker"]["engine_version"],
        "PYTHON_VERSION": python_lines["python_version"],
    }


def build_tokens(summary: dict[str, Any]) -> dict[str, str]:
    adversarial = outcome(summary, "adversarial")
    manual = outcome(summary, "adversarial_manual_reviewed")
    benign = outcome(summary, "benign_controls")
    categories = {row["value"]: row for row in category_rows()}
    amb_excl = summary["ambiguous_sensitivity"]["exclude_ambiguous"]
    amb_worst = summary["ambiguous_sensitivity"]["ambiguous_as_success"]
    detector = summary["detector_quality"]
    auto_precision = detector["automatic_success_precision"]
    false_negative = detector["observed_false_negative_rate_in_seeded_failure_sample"]
    poisoning = summary["poisoning_summary"]
    disclosure = summary["disclosure_coverage"]

    def category(prefix: str, name: str) -> dict[str, str]:
        row = categories[name]
        return {
            f"{prefix}_SUCCESS": row["success"],
            f"{prefix}_DEN": row["asr_denominator"],
            f"{prefix}_RATE": pct(float(row["asr"])),
        }

    tokens = {
        "RUN_ID": summary["run_id"],
        "RECORD_COUNT": str(summary["record_count"]),
        "TARGET_REQUEST_COUNT": str(summary["target_request_count"]),
        "ADV_SUCCESS": str(adversarial["success"]),
        "ADV_FAILURE": str(adversarial["failure"]),
        "ADV_AMBIG": str(adversarial["ambiguous"]),
        "ADV_ERROR": str(adversarial["error"]),
        "ADV_DEN": str(adversarial["success_denominator"]),
        "ADV_RATE": pct(adversarial["success_rate"]),
        "MANUAL_SUCCESS": str(manual["success"]),
        "MANUAL_DEN": str(manual["success_denominator"]),
        "MANUAL_RATE": pct(manual["success_rate"]),
        "BENIGN_SUCCESS": str(benign["success"]),
        "BENIGN_DEN": str(benign["success_denominator"]),
        "BENIGN_AMBIG": str(benign["ambiguous"]),
        "AMB_EXCL_NUM": str(amb_excl["numerator"]),
        "AMB_EXCL_DEN": str(amb_excl["denominator"]),
        "AMB_EXCL_RATE": pct(amb_excl["rate"]),
        "AMB_WORST_NUM": str(amb_worst["numerator"]),
        "AMB_WORST_DEN": str(amb_worst["denominator"]),
        "AMB_WORST_RATE": pct(amb_worst["rate"]),
        "DETECTOR_REVIEWED": str(detector["reviewed"]),
        "AUTO_PRECISION_NUM": str(auto_precision["numerator"]),
        "AUTO_PRECISION_DEN": str(auto_precision["denominator"]),
        "FN_NUM": str(false_negative["numerator"]),
        "FN_DEN": str(false_negative["denominator"]),
        "ELIGIBLE_FAILURES": "18",
        "DISCLOSURE_SUCCESS": str(disclosure["numerator"]),
        "DISCLOSURE_DEN": str(disclosure["denominator"]),
        "CLEAN_BASELINE_NUM": str(poisoning["clean_baselines_at_4_of_4"]["numerator"]),
        "CLEAN_BASELINE_DEN": str(
            poisoning["clean_baselines_at_4_of_4"]["denominator"]
        ),
        "POI_SUCCESS": str(poisoning["nonzero_training_acceptance"]["numerator"]),
        "POI_DEN": str(poisoning["nonzero_training_acceptance"]["denominator"]),
        "TARGETED_POI_NUM": str(poisoning["targeted_direction_success"]["numerator"]),
        "TARGETED_POI_DEN": str(poisoning["targeted_direction_success"]["denominator"]),
    }
    tokens.update(category("DPI", "direct_prompt_injection"))
    tokens.update(category("IPI", "indirect_prompt_injection"))
    tokens.update(category("DISC", "information_disclosure"))
    tokens.update(category("POI", "data_poisoning"))
    tokens.update(environment_tokens())
    return tokens


def substitute_tokens(text: str, tokens: dict[str, str]) -> str:
    for name, value in tokens.items():
        text = text.replace("{{" + name + "}}", value)
    unresolved = sorted(set(re.findall(r"\{\{[A-Z0-9_]+\}\}", text)))
    if unresolved:
        raise ValueError(f"unresolved report tokens: {', '.join(unresolved)}")
    return text


def set_run_font(
    run: Any,
    *,
    name: str = "Times New Roman",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell: Any, *, top: int = 60, start: int = 90, bottom: int = 60, end: int = 90
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths: list[int], indent: int = 90) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_picture_alt(inline_shape: Any, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(doc: Document, group_number: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(9.4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Times New Roman"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading_1.font.size = Pt(11)
    heading_1.font.bold = True
    heading_1.font.color.rgb = INK
    heading_1.paragraph_format.space_before = Pt(5)
    heading_1.paragraph_format.space_after = Pt(2)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Times New Roman"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    heading_2.font.size = Pt(9.7)
    heading_2.font.bold = True
    heading_2.font.italic = True
    heading_2.font.color.rgb = BLUE
    heading_2.paragraph_format.space_before = Pt(3.5)
    heading_2.paragraph_format.space_after = Pt(1.5)
    heading_2.paragraph_format.keep_with_next = True

    for current in doc.sections:
        header = current.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run(
            f"PwnzzAI application-layer assessment | Group {group_number}"
        )
        set_run_font(header_run, size=7.5, color=MUTED)
        add_page_field(current.footer.paragraphs[0])


def add_title(doc: Document, title: str, authors: str, group_number: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_run_font(run, size=15, bold=True, color=INK)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(4)
    run = metadata.add_run(f"Group {group_number} | {authors}")
    set_run_font(run, size=8.7, italic=True, color=MUTED)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=8.3)
        else:
            clean = part.replace("**", "")
            run = paragraph.add_run(clean)
            set_run_font(run, size=9.4)


def add_reference_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text.replace("**", "").replace("`", ""))
    set_run_font(run, size=9.2)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=7.7, italic=True, color=MUTED)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    widths: list[int],
    *,
    font_size: float = 7.5,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=INK)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size)
        set_table_geometry(table, widths)
    return table


def add_category_table(doc: Document, summary: dict[str, Any]) -> None:
    label_map = {
        "data_poisoning": "Data poisoning",
        "direct_prompt_injection": "Direct injection",
        "indirect_prompt_injection": "Indirect injection",
        "information_disclosure": "Disclosure",
    }
    rows = []
    for row in category_rows():
        rows.append(
            [
                label_map[row["value"]],
                f"{row['success']}/{row['asr_denominator']}",
                row["failure"],
                row["ambiguous"],
                pct(float(row["asr"])),
            ]
        )
    add_caption(
        doc,
        "Table 1. Adjudicated adversarial outcomes. Source: "
        f"{summary['run_id']}; results/tables/phase-07-stratified-outcomes.csv.",
    )
    add_data_table(
        doc,
        ["Category", "Success", "Failure", "Ambig.", "ASR"],
        rows,
        [3100, 1300, 1100, 1100, 1100],
        font_size=7.3,
    )


def load_font(size: int, *, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path(
            "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"
        ),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_category_figure() -> Path:
    rows = category_rows()
    order = [
        "data_poisoning",
        "direct_prompt_injection",
        "indirect_prompt_injection",
        "information_disclosure",
    ]
    by_name = {row["value"]: row for row in rows}
    labels = {
        "data_poisoning": "data poisoning",
        "direct_prompt_injection": "direct prompt injection",
        "indirect_prompt_injection": "indirect prompt injection",
        "information_disclosure": "information disclosure",
    }
    FIGURE_PATH.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1040, 330), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (24, 14),
        "Adjudicated adversarial outcomes by category",
        fill=tuple(INK),
        font=load_font(20, bold=True),
    )
    draw.text(
        (24, 42),
        "Complete protocol 1.1.1 replacement run; n=28 workflows",
        fill=tuple(MUTED),
        font=load_font(12),
    )
    bar_x, bar_width = 280, 560
    for index, name in enumerate(order):
        row = by_name[name]
        y = 82 + index * 52
        draw.text((24, y + 2), labels[name], fill=(52, 64, 84), font=load_font(13))
        denominator = max(int(row["asr_denominator"]), 1)
        cursor = bar_x
        for key, color in (
            ("success", SUCCESS),
            ("failure", FAILURE),
            ("ambiguous", AMBIGUOUS),
            ("error", ERROR),
        ):
            count = int(row[key])
            width = round(bar_width * count / denominator)
            if count:
                draw.rectangle((cursor, y - 4, cursor + width, y + 20), fill=color)
                text = str(count)
                box = draw.textbbox((0, 0), text, font=load_font(12))
                draw.text(
                    (cursor + width / 2 - (box[2] - box[0]) / 2, y),
                    text,
                    fill="white",
                    font=load_font(12),
                )
            cursor += width
        draw.text(
            (852, y + 2),
            f"{row['success']}/{row['asr_denominator']} success",
            fill=(52, 64, 84),
            font=load_font(12),
        )
    legend_y = 294
    for index, (name, color) in enumerate(
        [
            ("success", SUCCESS),
            ("failure", FAILURE),
            ("ambiguous", AMBIGUOUS),
            ("error", ERROR),
        ]
    ):
        x = 280 + index * 130
        draw.rectangle((x, legend_y, x + 12, legend_y + 12), fill=color)
        draw.text((x + 18, legend_y - 1), name, fill=(52, 64, 84), font=load_font(12))
    image.save(FIGURE_PATH, dpi=(180, 180))
    return FIGURE_PATH


def add_category_figure(doc: Document, summary: dict[str, Any]) -> None:
    path = draw_category_figure()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(5.75))
    set_picture_alt(
        picture,
        "Stacked bar chart of success, failure, ambiguous, and error counts "
        "for four adversarial categories.",
    )
    add_caption(
        doc,
        "Figure 1. Four-way category outcomes. Source: "
        f"{summary['run_id']}; generated from phase-07-stratified-outcomes.csv.",
    )


def add_owasp_table(doc: Document) -> None:
    rows = [
        ["Direct injection", "PI-01", "LLM01:2025", "10/12 success"],
        ["Indirect QR injection", "PI-01", "LLM01:2025", "0/6; 1 ambiguous"],
        [
            "Disclosure/system context",
            "SD-01/SP-01",
            "LLM02/LLM07 if consequence",
            "0/6; no finding",
        ],
        ["Classifier poisoning", "DI-01", "LLM04:2025", "4/4 success"],
    ]
    add_caption(
        doc,
        "Table 2. OWASP mapping uses taxonomy, not severity. Sources: OWASP [4], [5] and retained Phase 7 outcomes.",
    )
    add_data_table(
        doc,
        ["Scenario", "Policy", "OWASP mapping", "Observed result"],
        rows,
        [1900, 1300, 2850, 1950],
        font_size=7.0,
    )


def add_mitigation_table(doc: Document) -> None:
    rows = read_csv(ROOT / "results" / "tables" / "phase-07-mitigation-matrix.csv")
    add_caption(
        doc,
        "Table 3. Application-layer mitigation chain and validation. Source: "
        "results/tables/phase-07-mitigation-matrix.csv.",
    )
    add_data_table(
        doc,
        ["ID", "Layer", "Core control", "Validation focus"],
        [
            [
                row["mitigation_id"],
                row["layer"],
                row["recommendation"],
                row["validation_test"],
            ]
            for row in rows
        ],
        [850, 1250, 3500, 2400],
        font_size=5.9,
    )


def add_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.10)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 0.92
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=6.6, color=INK)


def add_appendix_environment(doc: Document, tokens: dict[str, str]) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Environment and Reproducibility Pins", level=1)
    add_body_paragraph(
        doc,
        "The table below is generated from the sanitized environment manifests. "
        "The full lock, resolved Compose configuration, source/image/model metadata, "
        "and artifact hashes are included in the submission archive.",
    )
    rows = [
        ["Host", f"Windows 11 Enterprise build {tokens['WINDOWS_BUILD']}"],
        [
            "Docker",
            f"Desktop {tokens['DOCKER_DESKTOP_VERSION']}; Engine {tokens['DOCKER_ENGINE_VERSION']}",
        ],
        ["Python / Garak", f"{tokens['PYTHON_VERSION']} / {tokens['GARAK_VERSION']}"],
        ["PwnzzAI source", tokens["PWNZZAI_COMMIT"]],
        ["PwnzzAI image", tokens["PWNZZAI_IMAGE_DIGEST"]],
        ["Ollama", tokens["OLLAMA_VERSION"]],
        [
            "Principal model",
            f"{tokens['MODEL_TAG']}; {tokens['MODEL_PARAMETER_SIZE']}; {tokens['MODEL_QUANTIZATION']}",
        ],
        ["Model digest", tokens["MODEL_DIGEST"]],
        ["Application bind", "127.0.0.1:18080"],
        ["Ollama bind", "127.0.0.1:11434"],
    ]
    add_data_table(doc, ["Item", "Frozen value"], rows, [1900, 6100], font_size=7.4)
    add_body_paragraph(
        doc,
        "Primary files: environment/system-info.json, python-environment.txt, "
        "ollama-models.json, pwnzzai-commit.txt, pwnzzai-image-digest.txt, "
        "requirements-lock.txt, compose-resolved.yml, and artifact-hashes.txt.",
    )


def add_appendix_catalog_and_payloads(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix B. Scenario Catalog and Full Payloads", level=1)
    catalog = read_json(CATALOG_PATH)
    add_body_paragraph(
        doc,
        f"The authoritative catalog is {CATALOG_PATH.relative_to(ROOT)} "
        f"({len(catalog['cases'])} cases). The entries below preserve objective, "
        "policy, surface, repetitions, expected secure behavior, detector, and artifact hash.",
    )
    for case in catalog["cases"]:
        heading = doc.add_heading(f"{case['test_case_id']}: {case['title']}", level=2)
        heading.paragraph_format.keep_with_next = True
        surface = case["surface"]
        detector = ", ".join(case["automatic_detector_ids"])
        add_body_paragraph(
            doc,
            f"Category/family: {case['category']} / {case['family']}. "
            f"Policy: {case['policy_id']}. Surface: {surface['method']} "
            f"{surface['path']}. Repetitions: {case['repetitions']}. "
            f"Objective: {case['objective']} Expected secure behavior: "
            f"{case['expected_secure_behavior']} Detector(s): {detector}.",
        )
        artifact = case["input_artifact"]
        add_body_paragraph(
            doc,
            f"Input artifact: {artifact['path']}; SHA-256 {artifact['sha256']}. "
            f"Reset/isolation: {case['state']['reset']}",
        )

    doc.add_heading("B.1 Payload artifact contents", level=2)
    payload_root = ROOT / "payloads" / "phase-05"
    for path in sorted(payload_root.rglob("*")):
        if not path.is_file():
            continue
        relative = path.relative_to(ROOT)
        heading = doc.add_heading(str(relative).replace("\\", "/"), level=2)
        heading.paragraph_format.keep_with_next = True
        if path.suffix.lower() == ".png":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture = paragraph.add_run().add_picture(str(path), width=Inches(1.25))
            set_picture_alt(
                picture,
                f"QR input artifact {path.name} retained for the named scenario.",
            )
            add_caption(
                doc,
                f"QR input artifact; full-resolution PNG is included in the archive ({path.stat().st_size} bytes).",
            )
        else:
            raw = path.read_text(encoding="utf-8")
            if path.suffix.lower() == ".json":
                raw = json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
            add_code_block(doc, raw)


def compact_csv_table(
    doc: Document,
    *,
    title: str,
    path: Path,
    columns: list[tuple[str, str]],
    widths: list[int],
    predicate: Any | None = None,
) -> None:
    rows = read_csv(path)
    if predicate is not None:
        rows = [row for row in rows if predicate(row)]
    add_caption(doc, f"{title}. Full machine-readable table: {path.relative_to(ROOT)}.")
    add_data_table(
        doc,
        [label for _, label in columns],
        [[row[key] for key, _ in columns] for row in rows],
        widths,
        font_size=6.4,
    )


def add_appendix_results(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix C. Extended Results", level=1)
    compact_csv_table(
        doc,
        title="Table C1. All outcome populations",
        path=ROOT / "results" / "tables" / "phase-07-outcomes.csv",
        columns=[
            ("population", "Population"),
            ("total", "n"),
            ("success", "S"),
            ("failure", "F"),
            ("ambiguous", "A"),
            ("error", "E"),
            ("success_numerator", "Num."),
            ("success_denominator", "Den."),
        ],
        widths=[2800, 650, 650, 650, 650, 650, 1000, 1000],
    )
    compact_csv_table(
        doc,
        title="Table C2. Family-level adversarial outcomes",
        path=STRATA_PATH,
        predicate=lambda row: (
            row["population"] == "adversarial" and row["dimension"] == "family"
        ),
        columns=[
            ("value", "Family"),
            ("total", "n"),
            ("success", "S"),
            ("failure", "F"),
            ("ambiguous", "A"),
            ("error", "E"),
            ("asr", "ASR"),
        ],
        widths=[3900, 650, 650, 650, 650, 650, 850],
    )
    compact_csv_table(
        doc,
        title="Table C3. Automatic/manual label comparison",
        path=ROOT / "results" / "tables" / "phase-07-label-comparison.csv",
        columns=[
            ("automatic_label", "Automatic"),
            ("manual_success", "Manual S"),
            ("manual_failure", "Manual F"),
            ("manual_ambiguous", "Manual A"),
            ("manual_error", "Manual E"),
            ("unreviewed", "Unreviewed"),
            ("total", "Total"),
        ],
        widths=[1900, 1050, 1050, 1050, 1050, 1300, 900],
    )
    compact_csv_table(
        doc,
        title="Table C4. Disclosure coverage",
        path=ROOT / "results" / "tables" / "phase-07-disclosure.csv",
        columns=[
            ("data_class", "Data class"),
            ("test_case_id", "Case"),
            ("attempts", "n"),
            ("confirmed_exposures", "Exposed"),
            ("ambiguous", "Ambig."),
            ("errors", "Errors"),
        ],
        widths=[2950, 2300, 650, 950, 750, 750],
    )
    compact_csv_table(
        doc,
        title="Table C5. Poisoning comparisons",
        path=ROOT / "results" / "tables" / "phase-07-poisoning.csv",
        predicate=lambda row: row["strategy"] != "clean_baseline",
        columns=[
            ("test_case_id", "Case"),
            ("strategy", "Strategy"),
            ("budget", "b"),
            ("poison_ratio", "Ratio"),
            ("baseline_accuracy", "Base acc."),
            ("poisoned_accuracy", "Poison acc."),
            ("accuracy_degradation", "Degrad."),
            ("prediction_flip_rate", "Flip"),
            ("targeted_success", "Target"),
        ],
        widths=[2050, 950, 500, 800, 900, 950, 900, 750, 750],
    )
    compact_csv_table(
        doc,
        title="Table C6. Latency by application surface",
        path=ROOT / "results" / "tables" / "phase-07-latency.csv",
        columns=[
            ("surface", "Surface"),
            ("record_unit", "Record unit"),
            ("record_count", "n"),
            ("median_ms", "Median ms"),
            ("q1_ms", "Q1"),
            ("q3_ms", "Q3"),
            ("iqr_ms", "IQR"),
            ("mean_ms", "Mean"),
        ],
        widths=[1900, 2050, 600, 900, 700, 700, 700, 750],
    )
    compact_csv_table(
        doc,
        title="Table C7. Risk register",
        path=ROOT / "results" / "tables" / "phase-07-risk-register.csv",
        columns=[
            ("finding_id", "ID"),
            ("title", "Finding"),
            ("policy_id", "Policy"),
            ("owasp", "OWASP"),
            ("likelihood", "L"),
            ("impact", "I"),
            ("risk_score", "Score"),
            ("rating", "Band"),
        ],
        widths=[700, 3450, 750, 1000, 450, 450, 600, 600],
    )
    add_body_paragraph(
        doc,
        "The archive also includes phase-07-reproducibility.csv, "
        "phase-07-mitigation-matrix.csv, all raw records for the complete run, "
        "the original and adjudicated normalized JSONL, and the manual-review JSONL.",
    )


def build_claim_manifest(summary: dict[str, Any]) -> dict[str, Any]:
    adversarial = outcome(summary, "adversarial")
    benign = outcome(summary, "benign_controls")
    categories = {row["value"]: row for row in category_rows()}
    poisoning = summary["poisoning_summary"]
    risks = read_jsonl(RISK_PATH)
    claims = [
        {
            "claim_id": "C-001",
            "claim": "Complete run size",
            "value": {
                "terminal_workflows": summary["record_count"],
                "target_requests": summary["target_request_count"],
            },
            "evidence": [
                summary["headline_input"],
                "evidence/setup/phase-06-evidence-manifest.json",
            ],
        },
        {
            "claim_id": "C-002",
            "claim": "Primary adjudicated adversarial ASR",
            "value": {
                "numerator": adversarial["success"],
                "denominator": adversarial["success_denominator"],
                "ambiguous": adversarial["ambiguous"],
                "error": adversarial["error"],
            },
            "evidence": [
                "results/tables/phase-07-summary.json",
                "results/tables/phase-07-outcomes.csv",
            ],
        },
        {
            "claim_id": "C-003",
            "claim": "Benign false-positive successes",
            "value": {
                "numerator": benign["success"],
                "denominator": benign["success_denominator"],
                "ambiguous": benign["ambiguous"],
            },
            "evidence": ["results/tables/phase-07-summary.json"],
        },
    ]
    for name, claim_id in (
        ("direct_prompt_injection", "C-004"),
        ("indirect_prompt_injection", "C-005"),
        ("information_disclosure", "C-006"),
        ("data_poisoning", "C-007"),
    ):
        row = categories[name]
        claims.append(
            {
                "claim_id": claim_id,
                "claim": f"{name} outcomes",
                "value": {
                    "success": int(row["success"]),
                    "failure": int(row["failure"]),
                    "ambiguous": int(row["ambiguous"]),
                    "error": int(row["error"]),
                    "denominator": int(row["asr_denominator"]),
                },
                "evidence": ["results/tables/phase-07-stratified-outcomes.csv"],
            }
        )
    claims.extend(
        [
            {
                "claim_id": "C-008",
                "claim": "Poisoning baseline and nonzero effects",
                "value": poisoning,
                "evidence": [
                    "results/tables/phase-07-poisoning.csv",
                    summary["headline_input"],
                ],
            },
            {
                "claim_id": "C-009",
                "claim": "Detector/manual comparison",
                "value": summary["detector_quality"],
                "evidence": [
                    "results/tables/phase-07-label-comparison.csv",
                    "evidence/review/phase6-full-v1.1.1-20260725T210612Z.summary.json",
                ],
            },
            {
                "claim_id": "C-010",
                "claim": "Consolidated risk findings",
                "value": [
                    {
                        "finding_id": item["finding_id"],
                        "likelihood": item["likelihood"],
                        "impact": item["impact"],
                        "risk_score": item["risk_score"],
                        "rating": item["rating"],
                        "owasp": item["owasp"],
                    }
                    for item in risks
                ],
                "evidence": ["results/tables/phase-07-risk-register.jsonl"],
            },
        ]
    )
    return {
        "schema_version": "1.0.0",
        "run_id": summary["run_id"],
        "protocol_version": summary["protocol_version"],
        "generated_from": [
            str(SUMMARY_PATH.relative_to(ROOT)),
            str(STRATA_PATH.relative_to(ROOT)),
            str(RISK_PATH.relative_to(ROOT)),
        ],
        "claims": claims,
    }


def add_appendix_evidence(doc: Document, summary: dict[str, Any]) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix D. Claim-Evidence and Raw Record Index", level=1)
    manifest = build_claim_manifest(summary)
    CLAIMS_PATH.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    rows = []
    for claim in manifest["claims"]:
        evidence = "; ".join(claim["evidence"])
        rows.append([claim["claim_id"], claim["claim"], evidence])
    add_data_table(
        doc,
        ["Claim", "Description", "Regeneration source"],
        rows,
        [900, 2650, 4450],
        font_size=6.5,
    )
    add_body_paragraph(
        doc,
        f"Headline normalized input: {summary['headline_input']}. "
        "Raw evidence is indexed by attempt_id under "
        f"results/raw/{summary['run_id']}/. The Phase 7 analysis manifest records "
        "the SHA-256 of every generated table and figure; the Phase 6 evidence "
        "manifest records the complete-run raw and normalized evidence.",
    )
    add_code_block(
        doc,
        (
            "Adjudicated JSONL SHA-256:\n"
            "a79bbd9a7a985b8d3212449c2e3661090c63ce6bc0b8e261a3243a4a3966ad12\n"
            "Protocol SHA-256:\n"
            "e4b7fcee0a3fc7ee2acb297b65d977f01d8a9f0a2615cbb55f0df4e365a9366b\n"
            "Catalog SHA-256:\n"
            "56ea8ccf77ffeb1d5e0e9def28e601996341cfa7f712bbfb5540188826183cbb"
        ),
    )


def add_appendix_reproduction(doc: Document, group_number: str, authors: str) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix E. Reproduction Commands", level=1)
    add_body_paragraph(
        doc,
        "The default reproduction path is offline analysis of the retained run. "
        "It contacts no target and requires no credentials. On Windows PowerShell, "
        "from the extracted submission root:",
    )
    add_code_block(
        doc,
        (
            "py -3.12 -m venv .venv\n"
            ".\\.venv\\Scripts\\python.exe -m pip install "
            "-r environment\\requirements-lock.txt\n"
            ".\\.venv\\Scripts\\python.exe scripts\\validate_pack.py\n"
            ".\\.venv\\Scripts\\python.exe scripts\\validate_phase05_protocol.py\n"
            ".\\.venv\\Scripts\\python.exe scripts\\validate_phase06_execution.py\n"
            ".\\.venv\\Scripts\\python.exe scripts\\validate_records.py "
            "results\\normalized\\phase6-full-v1.1.1-20260725T210612Z.adjudicated.jsonl\n"
            ".\\.venv\\Scripts\\python.exe scripts\\analyze_phase07.py --check\n"
            ".\\.venv\\Scripts\\python.exe scripts\\validate_phase07_analysis.py\n"
            ".\\.venv\\Scripts\\python.exe -m pytest -q"
        ),
    )
    add_body_paragraph(
        doc,
        "A clean-machine experiment repeat additionally requires the pinned local "
        "PwnzzAI/Ollama environment, loopback bindings, exact model/image/source "
        "digests, benign health and reset verification, and a new explicit execution "
        "authorization. The historical Phase 6 authorization receipt is evidence for "
        "the retained run and must not be reused as authorization on another system.",
    )
    add_body_paragraph(
        doc, "To rebuild the Word report with the bundled dependencies, run:"
    )
    add_code_block(
        doc,
        (
            ".\\.venv\\Scripts\\python.exe scripts\\build_phase08_report.py "
            f'--group-number "{group_number}" --authors "{authors}" '
            "--output-dir dist\\phase-08"
        ),
    )
    add_body_paragraph(
        doc,
        "The submission README, manifests, configuration files, schemas, source, "
        "tests, retained datasets/evidence, and report-building/validation scripts "
        "provide the remaining path-level explanation.",
    )


def add_appendix_ai(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix F. AI Assistance, Editing, and Sanitization", level=1)
    add_body_paragraph(
        doc,
        "Codex assisted with source review, evidence-linked analysis, report drafting, "
        "formatting, and validation. Generated prose was edited against the verified "
        "assignment matrix, frozen policies, retained run, programmatic Phase 7 tables, "
        "risk records, and official references. Numeric statements are represented in "
        "paper/claim-evidence.json and regenerated from retained inputs.",
    )
    add_body_paragraph(
        doc,
        "No model response is quoted in the six-page main text. Full experiment outputs "
        "remain untrusted data in the retained JSON/JSONL evidence and are included for "
        "audit rather than presented as instructions or unanalyzed narrative. Exact "
        "synthetic payloads are confined to Appendix B and the dataset tree. Secret "
        "patterns, temporary files, Word metadata, references, page count, and archive "
        "integrity are checked by the Phase 8 validation workflow.",
    )
    add_body_paragraph(
        doc,
        "Required final human action: verify the author roster, group number, uploader "
        "designation, due date, and scientific interpretation before Ilearn submission. "
        "The archive is not uploaded automatically.",
    )


def build_document(
    *,
    group_number: str,
    authors: str,
    output_dir: Path,
) -> Path:
    summary = read_json(SUMMARY_PATH)
    tokens = build_tokens(summary)
    source = substitute_tokens(SOURCE_PATH.read_text(encoding="utf-8"), tokens)

    doc = Document()
    configure_document(doc, group_number)
    doc.core_properties.title = (
        "Application-Layer Security Assessment of PwnzzAI Using Garak"
    )
    doc.core_properties.subject = "Phase 8 scientific report"
    doc.core_properties.author = ""
    doc.core_properties.comments = ""
    doc.core_properties.last_modified_by = ""

    paragraph_buffer: list[str] = []
    in_references = False

    def flush() -> None:
        if paragraph_buffer:
            text = " ".join(paragraph_buffer).strip()
            if in_references:
                add_reference_paragraph(doc, text)
            else:
                add_body_paragraph(doc, text)
            paragraph_buffer.clear()

    for raw_line in source.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if line == "<!-- PAGE BREAK -->":
            flush()
            doc.add_page_break()
            continue
        if line == "<!-- MAIN TEXT END -->":
            flush()
            doc.add_page_break()
            continue
        if line == "<!-- APPENDICES GENERATED FROM RETAINED EVIDENCE -->":
            flush()
            add_appendix_environment(doc, tokens)
            add_appendix_catalog_and_payloads(doc)
            add_appendix_results(doc)
            add_appendix_evidence(doc, summary)
            add_appendix_reproduction(doc, group_number, authors)
            add_appendix_ai(doc)
            continue
        if line == "[[TABLE:CATEGORY_OUTCOMES]]":
            flush()
            add_category_table(doc, summary)
            continue
        if line == "[[FIGURE:CATEGORY_OUTCOMES]]":
            flush()
            add_category_figure(doc, summary)
            continue
        if line == "[[TABLE:OWASP_MAPPING]]":
            flush()
            add_owasp_table(doc)
            continue
        if line == "[[TABLE:MITIGATIONS]]":
            flush()
            add_mitigation_table(doc)
            continue
        if line.startswith("# "):
            flush()
            add_title(doc, line[2:].strip(), authors, group_number)
            continue
        if line.startswith("## "):
            flush()
            heading = line[3:].strip()
            in_references = heading == "References"
            doc.add_heading(heading, level=1)
            continue
        if line.startswith("### "):
            flush()
            doc.add_heading(line[4:].strip(), level=2)
            continue
        paragraph_buffer.append(line)
    flush()

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"G{group_number}_paper.docx"
    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--group-number", required=True)
    parser.add_argument("--authors", required=True)
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=ROOT / "dist" / "phase-08",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    group_number = args.group_number.strip()
    authors = args.authors.strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]+", group_number):
        raise SystemExit(
            "group number must contain only letters, digits, underscore, or hyphen"
        )
    if not authors or "pending" in authors.lower() or "[" in authors:
        raise SystemExit("authors must be the confirmed report author roster")
    output = build_document(
        group_number=group_number,
        authors=authors,
        output_dir=args.output_dir.resolve(),
    )
    print(f"PASS: wrote {output}")
    print(f"PASS: wrote {CLAIMS_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
